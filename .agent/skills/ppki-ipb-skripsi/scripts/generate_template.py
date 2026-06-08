import os
import urllib.request
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_page_number(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def download_ipb_logo():
    logo_path = os.path.join(os.path.dirname(__file__), 'ipb_logo.png')
    if not os.path.exists(logo_path):
        try:
            url = "https://upload.wikimedia.org/wikipedia/id/0/0f/Logo_IPB.png"
            req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
            with urllib.request.urlopen(req) as response, open(logo_path, 'wb') as out_file:
                out_file.write(response.read())
        except Exception as e:
            print(f"Could not download logo: {e}")
    return logo_path

def create_ipb_template(output_path):
    logo_path = download_ipb_logo()
    doc = Document()

    # Enable mirror margins in settings
    settings = doc.settings
    mirror_margins = OxmlElement('w:mirrorMargins')
    settings._element.append(mirror_margins)
    settings.odd_and_even_pages_header_footer = True

    def setup_section(section, add_headers=True):
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(3)
        section.left_margin = Cm(4) # Inside margin (pias kiri)
        section.right_margin = Cm(3) # Outside margin (batas kanan)
        section.page_height = Cm(29.7)
        section.page_width = Cm(21)
        section.header_distance = Cm(2) # batas atas 2 cm untuk nomor halaman
        section.footer_distance = Cm(2)

        if add_headers:
            header_odd = section.header
            header_odd.is_linked_to_previous = False
            if len(header_odd.paragraphs) == 0:
                p_odd = header_odd.add_paragraph()
            else:
                p_odd = header_odd.paragraphs[0]
            p_odd.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_odd.style.font.name = 'Times New Roman'
            p_odd.style.font.size = Pt(12)
            p_odd.clear()
            run_odd = p_odd.add_run()
            run_odd.font.name = 'Times New Roman'
            run_odd.font.size = Pt(12)
            create_page_number(run_odd)

            header_even = section.even_page_header
            header_even.is_linked_to_previous = False
            if len(header_even.paragraphs) == 0:
                p_even = header_even.add_paragraph()
            else:
                p_even = header_even.paragraphs[0]
            p_even.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_even.style.font.name = 'Times New Roman'
            p_even.style.font.size = Pt(12)
            p_even.clear()
            run_even = p_even.add_run()
            run_even.font.name = 'Times New Roman'
            run_even.font.size = Pt(12)
            create_page_number(run_even)
        else:
            section.header.is_linked_to_previous = False
            section.even_page_header.is_linked_to_previous = False
            for p in section.header.paragraphs: p.clear()
            for p in section.even_page_header.paragraphs: p.clear()

    # 2. Set Default Style (Times New Roman 12)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    p_format = doc.styles['Normal'].paragraph_format
    p_format.line_spacing = 1.0 # 1 spasi
    p_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY # rata kanan (justified)
    p_format.first_line_indent = Cm(1) # menjorok 1 cm
    p_format.space_before = Pt(0)
    p_format.space_after = Pt(0)

    def add_dummy_top_paragraph():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(1)
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run()
        run.font.size = Pt(1)

    # --- SECTION 0: COVER AND TITLE PAGE (No page numbers) ---
    section0 = doc.sections[0]
    setup_section(section0, add_headers=False)

    # 1. Halaman Sampul (Cover - Lampiran 1b)
    add_dummy_top_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Cm(2) # Top Edge = 5cm. (3cm margin + 2cm spacing)
    run = p.add_run("JUDUL SKRIPSI\n(MAKSIMAL 3 BARIS, SPASI SATU, POSISI CENTER)")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Cm(6) # Top Edge = 12cm.
    run = p.add_run("NAMA LENGKAP MAHASISWA\nNIM")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Cm(4) # Top Edge = 18cm.
    run = p.add_run()
    if os.path.exists(logo_path):
        run.add_picture(logo_path, width=Cm(2.5))
    else:
        run.add_text("[LOGO IPB 2.5 CM]")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Cm(3.2) # Top Edge = 23.7cm.
    run = p.add_run("NAMA DEPARTEMEN/PROGRAM STUDI\nFAKULTAS/SEKOLAH\nINSTITUT PERTANIAN BOGOR\nBOGOR\n20XX")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True

    doc.add_page_break()

    # 2. Halaman Judul (Title Page - Lampiran 7b)
    add_dummy_top_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Cm(2) # Top Edge = 5cm.
    run = p.add_run("JUDUL SKRIPSI\n(MAKSIMAL 3 BARIS, SPASI SATU, POSISI CENTER)")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Cm(6) # Top Edge = 12cm.
    run = p.add_run("NAMA LENGKAP MAHASISWA\nNIM")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Cm(3) # Top Edge ~ 17cm.
    run = p.add_run("Skripsi\nsebagai salah satu syarat untuk memperoleh gelar\nSarjana pada\nProgram Studi .......")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = False

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Cm(4) # Top Edge ~ 23.7cm.
    run = p.add_run("NAMA DEPARTEMEN/PROGRAM STUDI\nFAKULTAS/SEKOLAH\nINSTITUT PERTANIAN BOGOR\nBOGOR\n20XX")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True


    # --- SECTION 1: FRONT MATTER (Roman Numerals) ---
    section1 = doc.add_section(WD_SECTION.NEW_PAGE)
    setup_section(section1, add_headers=True)
    
    def set_page_numbering(section, fmt="lowerRoman", start=None):
        sectPr = section._sectPr
        pgNumType = OxmlElement('w:pgNumType')
        pgNumType.set(qn('w:fmt'), fmt)
        if start is not None:
            pgNumType.set(qn('w:start'), str(start))
        sectPr.append(pgNumType)

    set_page_numbering(section1, fmt="lowerRoman", start=1)

    def add_chapter_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(24) # 2 spasi underneath
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'

    front_matter = [
        "HALAMAN PERNYATAAN",
        "ABSTRAK",
        "ABSTRACT",
        "RINGKASAN",
        "PRAKATA",
        "DAFTAR ISI",
        "DAFTAR TABEL",
        "DAFTAR GAMBAR",
        "DAFTAR LAMPIRAN"
    ]

    for title in front_matter:
        add_chapter_title(title)
        p = doc.add_paragraph(f"Uraian untuk {title} diletakkan di sini. Paragraf pertama menjorok 1 cm.")
        if title != front_matter[-1]:
            doc.add_page_break()

    # --- SECTION 2: MAIN CONTENT (Arabic Numerals) ---
    main_section = doc.add_section(WD_SECTION.NEW_PAGE)
    setup_section(main_section, add_headers=True)
    set_page_numbering(main_section, fmt="decimal", start=1)
    
    def add_sub_chapter(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(24) # berjarak 2 spasi dari atas
        p.paragraph_format.space_after = Pt(12)  # 1 spasi bawah
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'

    def add_sub_sub_chapter(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(18) # berjarak 1.5 spasi dari atas
        p.paragraph_format.space_after = Pt(12)  # 1 spasi bawah
        run = p.add_run(text)
        run.bold = False # tidak tebal
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'

    def add_list_item(text, level=1):
        p = doc.add_paragraph(text)
        p.paragraph_format.first_line_indent = Cm(0)
        if level == 1:
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.first_line_indent = Cm(-0.5)
        elif level == 2:
            p.paragraph_format.left_indent = Cm(1.5)
            p.paragraph_format.first_line_indent = Cm(-0.5)
        elif level == 3:
            p.paragraph_format.left_indent = Cm(2.0)
            p.paragraph_format.first_line_indent = Cm(-0.5)

    # Mimic Lampiran 17
    add_chapter_title("II TINJAUAN PUSTAKA")
    
    add_sub_chapter("2.1 Judul Subbab (Kata dalam judul diawali dengan huruf kapital dan dicetak tebal)")
    doc.add_paragraph("Uraian dengan deskripsi .......................................................................................................................................................................................................................................")
    
    add_list_item("a   ...", level=1)
    add_list_item("b   ...", level=1)
    add_list_item("c   ...", level=1)
    add_list_item("1)  ...", level=2)
    add_list_item("2)  ...", level=2)
    add_list_item("3)  ...", level=2)

    add_sub_sub_chapter("2.1.1 Judul Sub-subbab (Kata dalam judul diawali dengan huruf kapital dan dicetak tidak tebal)")
    doc.add_paragraph("Uraian dengan deskripsi .......................................................................................................................................................................................................................................")

    add_list_item("a)  ...", level=1)
    add_list_item("b)  ...", level=1)
    add_list_item("c)  ...", level=1)
    add_list_item("(1) ...", level=2)
    add_list_item("(2) ...", level=2)
    add_list_item("(3) ...", level=2)

    add_sub_sub_chapter("2.1.2 Judul Sub-subbab")
    
    add_sub_chapter("2.2 Judul Subbab")
    doc.add_paragraph("Uraian dengan deskripsi .......................................................................................................................................................................................................................................")

    add_list_item("a   ...", level=1)
    add_list_item("b   ...", level=1)
    add_list_item("c   ...", level=1)
    add_list_item("1)  ...", level=2)
    add_list_item("2)  ...", level=2)
    add_list_item("3)  ...", level=2)

    add_sub_sub_chapter("2.2.1 Judul Sub-subbab")
    doc.add_paragraph("Uraian dengan deskripsi .......................................................................................................................................................................................................................................")

    add_list_item("a)  ...", level=1)
    add_list_item("b)  ...", level=1)
    add_list_item("c)  ...", level=1)
    add_list_item("(1) ...", level=2)
    add_list_item("(2) ...", level=2)
    add_list_item("(3) ...", level=2)

    add_sub_sub_chapter("2.2.2 Judul Sub-subbab")
    add_sub_sub_chapter("2.2.3 Judul Sub-subbab")

    p = doc.add_paragraph("2.2.3.1 Judul sub-sub-subbab* (Hanya kata awal dalam judul diawali dengan huruf kapital dan dicetak tidak tebal)")
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)

    p = doc.add_paragraph("* Pengebaban 4 tingkat tidak disarankan pada penulisan Tugas Akhir")
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(1.0)

    # DAFTAR PUSTAKA
    ref_section = doc.add_section(WD_SECTION.NEW_PAGE)
    setup_section(ref_section, add_headers=True)
    add_chapter_title("DAFTAR PUSTAKA")
    p = doc.add_paragraph("Naim A. 2020. Judul. Jurnal. 1(1):1-10.")
    p.paragraph_format.first_line_indent = Cm(-1)
    p.paragraph_format.left_indent = Cm(1)

    # Save the document
    doc.save(output_path)
    print(f"Template successfully created at: {output_path}")

if __name__ == "__main__":
    target_file = os.path.join(os.getcwd(), "PPKI_IPB_Template.docx")
    create_ipb_template(target_file)
